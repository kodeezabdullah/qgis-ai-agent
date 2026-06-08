try:
    from qgis.core import (
        QgsProject,
        QgsExpression,
        QgsVectorLayer,
        QgsRasterLayer,
        QgsField,
        QgsCoordinateReferenceSystem,
        QgsCategorizedSymbolRenderer,
        QgsRendererCategory,
        QgsGraduatedSymbolRenderer,
        QgsSymbol,
        QgsClassificationQuantile,
        QgsVectorFileWriter,
        QgsHeatmapRenderer,
        QgsSingleSymbolRenderer,
    )
    from qgis.PyQt.QtCore import QVariant
    from qgis.PyQt.QtGui import QColor
    from qgis.utils import iface
    import processing
    import random
    QGIS_AVAILABLE = True
except ImportError:
    QGIS_AVAILABLE = False
    iface = None


def _find_layer_by_name(name):
    layers = QgsProject.instance().mapLayersByName(name)
    if not layers:
        return None
    return layers[0]


def get_layers():
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    result = []
    for layer_id, layer in QgsProject.instance().mapLayers().items():
        result.append({
            "name": layer.name(),
            "type": layer.type().name if hasattr(layer.type(), "name") else str(layer.type()),
            "id": layer_id,
        })
    return result


def get_attributes(layer_name):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    if not hasattr(layer, "fields"):
        return {"error": f"Layer '{layer_name}' has no attribute fields"}
    return list(layer.fields().names())


def run_query(layer_name, expression):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}

    expr = QgsExpression(expression)
    if expr.hasParserError():
        return {"error": f"Invalid expression: {expr.parserErrorString()}"}

    if not layer.setSubsetString(expression):
        return {"error": "Failed to apply filter (layer may not support subset strings)"}
    return {"success": True, "message": f"Filter applied to '{layer_name}'", "count": layer.featureCount()}


def buffer_analysis(layer_name, distance):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}

    try:
        params = {
            "INPUT": layer,
            "DISTANCE": distance,
            "SEGMENTS": 5,
            "END_CAP_STYLE": 0,
            "JOIN_STYLE": 0,
            "MITER_LIMIT": 2,
            "DISSOLVE": False,
            "OUTPUT": "memory:",
        }
        output = processing.run("native:buffer", params)
        new_layer = output["OUTPUT"]
        new_name = f"{layer_name}_buffer_{distance}"
        new_layer.setName(new_name)
        QgsProject.instance().addMapLayer(new_layer)
        return {"success": True, "layer_name": new_name}
    except Exception as e:
        return {"error": f"Buffer analysis failed: {e}"}


def highlight_features(layer_name, expression):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}

    expr = QgsExpression(expression)
    if expr.hasParserError():
        return {"error": f"Invalid expression: {expr.parserErrorString()}"}

    layer.selectByExpression(expression)
    count = layer.selectedFeatureCount()
    return {"success": True, "count": count}


def remove_layer(layer_name):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    QgsProject.instance().removeMapLayer(layer.id())
    return {"success": True, "message": f"Removed layer '{layer_name}'"}


def rename_layer(old_name, new_name):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(old_name)
    if layer is None:
        return {"error": f"Layer '{old_name}' not found"}
    layer.setName(new_name)
    return {"success": True, "message": f"Renamed '{old_name}' to '{new_name}'"}


def zoom_to_layer(layer_name):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    if iface is None:
        return {"error": "QGIS interface not available"}
    canvas = iface.mapCanvas()
    canvas.setExtent(layer.extent())
    canvas.refresh()
    return {"success": True, "message": f"Zoomed to '{layer_name}'"}


def toggle_visibility(layer_name, visible=True):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    node = QgsProject.instance().layerTreeRoot().findLayer(layer.id())
    if node is None:
        return {"error": f"Layer tree node for '{layer_name}' not found"}
    node.setItemVisibilityChecked(bool(visible))
    return {"success": True, "message": f"Layer '{layer_name}' visibility set to {bool(visible)}"}


def add_layer_from_path(file_path):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}

    import os
    if not os.path.exists(file_path):
        return {"error": f"File not found: {file_path}"}

    name = os.path.splitext(os.path.basename(file_path))[0]
    ext = os.path.splitext(file_path)[1].lower()
    raster_exts = {".tif", ".tiff", ".geotiff", ".img", ".asc", ".jp2", ".png", ".jpg", ".jpeg"}

    if ext in raster_exts:
        layer = QgsRasterLayer(file_path, name)
    else:
        layer = QgsVectorLayer(file_path, name, "ogr")

    if not layer.isValid():
        return {"error": f"Invalid layer: {file_path}"}

    QgsProject.instance().addMapLayer(layer)
    return {"success": True, "layer_name": layer.name()}


def _run_and_add(alg, params, new_name):
    try:
        output = processing.run(alg, params)
        new_layer = output["OUTPUT"]
        if hasattr(new_layer, "setName"):
            new_layer.setName(new_name)
            QgsProject.instance().addMapLayer(new_layer)
        return {"success": True, "layer_name": new_name}
    except Exception as e:
        return {"error": f"{alg} failed: {e}"}


def clip_layer(input_layer, mask_layer):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    a = _find_layer_by_name(input_layer)
    b = _find_layer_by_name(mask_layer)
    if a is None:
        return {"error": f"Layer '{input_layer}' not found"}
    if b is None:
        return {"error": f"Layer '{mask_layer}' not found"}
    return _run_and_add(
        "native:clip",
        {"INPUT": a, "OVERLAY": b, "OUTPUT": "memory:"},
        f"{input_layer}_clipped",
    )


def intersect_layers(layer1, layer2):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    a = _find_layer_by_name(layer1)
    b = _find_layer_by_name(layer2)
    if a is None:
        return {"error": f"Layer '{layer1}' not found"}
    if b is None:
        return {"error": f"Layer '{layer2}' not found"}
    return _run_and_add(
        "native:intersection",
        {"INPUT": a, "OVERLAY": b, "OUTPUT": "memory:"},
        f"{layer1}_intersect_{layer2}",
    )


def union_layers(layer1, layer2):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    a = _find_layer_by_name(layer1)
    b = _find_layer_by_name(layer2)
    if a is None:
        return {"error": f"Layer '{layer1}' not found"}
    if b is None:
        return {"error": f"Layer '{layer2}' not found"}
    return _run_and_add(
        "native:union",
        {"INPUT": a, "OVERLAY": b, "OUTPUT": "memory:"},
        f"{layer1}_union_{layer2}",
    )


def dissolve_layer(layer_name, field=None):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    params = {
        "INPUT": layer,
        "FIELD": [field] if field else [],
        "OUTPUT": "memory:",
    }
    return _run_and_add("native:dissolve", params, f"{layer_name}_dissolved")


def merge_layers(layer_names_list):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    if not isinstance(layer_names_list, list) or len(layer_names_list) < 2:
        return {"error": "Provide a list of at least 2 layer names"}
    layers = []
    for name in layer_names_list:
        layer = _find_layer_by_name(name)
        if layer is None:
            return {"error": f"Layer '{name}' not found"}
        layers.append(layer)
    return _run_and_add(
        "native:mergevectorlayers",
        {"LAYERS": layers, "OUTPUT": "memory:"},
        "merged_layer",
    )


def spatial_join(target_layer, join_layer):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    a = _find_layer_by_name(target_layer)
    b = _find_layer_by_name(join_layer)
    if a is None:
        return {"error": f"Layer '{target_layer}' not found"}
    if b is None:
        return {"error": f"Layer '{join_layer}' not found"}
    return _run_and_add(
        "native:joinattributesbylocation",
        {
            "INPUT": a,
            "JOIN": b,
            "PREDICATE": [0],
            "JOIN_FIELDS": [],
            "METHOD": 1,
            "DISCARD_NONMATCHING": False,
            "PREFIX": "",
            "OUTPUT": "memory:",
        },
        f"{target_layer}_joined",
    )


def _temp_raster_path(suffix=".tif"):
    import tempfile
    import os
    fd, path = tempfile.mkstemp(suffix=suffix, prefix="qgis_ai_agent_")
    os.close(fd)
    return path


def _run_raster_and_add(alg, params, new_name, output_key="OUTPUT"):
    try:
        out_path = _temp_raster_path()
        params[output_key] = out_path
        processing.run(alg, params)
        layer = QgsRasterLayer(out_path, new_name)
        if not layer.isValid():
            return {"error": f"{alg} produced invalid raster"}
        QgsProject.instance().addMapLayer(layer)
        return {"success": True, "layer_name": new_name, "path": out_path}
    except Exception as e:
        return {"error": f"{alg} failed: {e}"}


def calculate_ndvi(red_layer, nir_layer):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    red = _find_layer_by_name(red_layer)
    nir = _find_layer_by_name(nir_layer)
    if red is None:
        return {"error": f"Layer '{red_layer}' not found"}
    if nir is None:
        return {"error": f"Layer '{nir_layer}' not found"}
    return _run_raster_and_add(
        "gdal:rastercalculator",
        {
            "INPUT_A": nir,
            "BAND_A": 1,
            "INPUT_B": red,
            "BAND_B": 1,
            "FORMULA": "(A.astype(float)-B)/(A.astype(float)+B)",
            "NO_DATA": None,
            "RTYPE": 5,
        },
        "NDVI",
    )


def calculate_ndwi(green_layer, nir_layer):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    green = _find_layer_by_name(green_layer)
    nir = _find_layer_by_name(nir_layer)
    if green is None:
        return {"error": f"Layer '{green_layer}' not found"}
    if nir is None:
        return {"error": f"Layer '{nir_layer}' not found"}
    return _run_raster_and_add(
        "gdal:rastercalculator",
        {
            "INPUT_A": green,
            "BAND_A": 1,
            "INPUT_B": nir,
            "BAND_B": 1,
            "FORMULA": "(A.astype(float)-B)/(A.astype(float)+B)",
            "NO_DATA": None,
            "RTYPE": 5,
        },
        "NDWI",
    )


def hillshade(dem_layer):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    dem = _find_layer_by_name(dem_layer)
    if dem is None:
        return {"error": f"Layer '{dem_layer}' not found"}
    return _run_raster_and_add(
        "native:hillshade",
        {"INPUT": dem, "Z_FACTOR": 1, "AZIMUTH": 315, "V_ANGLE": 45},
        f"{dem_layer}_hillshade",
    )


def slope_analysis(dem_layer):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    dem = _find_layer_by_name(dem_layer)
    if dem is None:
        return {"error": f"Layer '{dem_layer}' not found"}
    return _run_raster_and_add(
        "native:slope",
        {"INPUT": dem, "Z_FACTOR": 1},
        f"{dem_layer}_slope",
    )


def aspect_analysis(dem_layer):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    dem = _find_layer_by_name(dem_layer)
    if dem is None:
        return {"error": f"Layer '{dem_layer}' not found"}
    return _run_raster_and_add(
        "native:aspect",
        {"INPUT": dem, "Z_FACTOR": 1},
        f"{dem_layer}_aspect",
    )


def zonal_statistics(raster_layer, vector_layer):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    raster = _find_layer_by_name(raster_layer)
    vector = _find_layer_by_name(vector_layer)
    if raster is None:
        return {"error": f"Layer '{raster_layer}' not found"}
    if vector is None:
        return {"error": f"Layer '{vector_layer}' not found"}
    try:
        output = processing.run(
            "native:zonalstatisticsfb",
            {
                "INPUT": vector,
                "INPUT_RASTER": raster,
                "RASTER_BAND": 1,
                "COLUMN_PREFIX": "_",
                "STATISTICS": [0, 1, 2, 3, 4, 5, 6, 7],
                "OUTPUT": "memory:",
            },
        )
        new_layer = output["OUTPUT"]
        new_name = f"{vector_layer}_zonalstats"
        new_layer.setName(new_name)
        QgsProject.instance().addMapLayer(new_layer)
        return {"success": True, "layer_name": new_name}
    except Exception as e:
        return {"error": f"zonal_statistics failed: {e}"}


def raster_clip(raster_layer, mask_layer):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    raster = _find_layer_by_name(raster_layer)
    mask = _find_layer_by_name(mask_layer)
    if raster is None:
        return {"error": f"Layer '{raster_layer}' not found"}
    if mask is None:
        return {"error": f"Layer '{mask_layer}' not found"}
    return _run_raster_and_add(
        "gdal:cliprasterbymasklayer",
        {
            "INPUT": raster,
            "MASK": mask,
            "CROP_TO_CUTLINE": True,
            "KEEP_RESOLUTION": True,
        },
        f"{raster_layer}_clipped",
    )


def unsupervised_classification(layer_name, num_classes):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    raster = _find_layer_by_name(layer_name)
    if raster is None:
        return {"error": f"Layer '{layer_name}' not found"}
    try:
        n = int(num_classes)
    except (TypeError, ValueError):
        return {"error": f"num_classes must be an integer, got {num_classes!r}"}
    return _run_raster_and_add(
        "saga:kmeansclusteringforgrids",
        {
            "GRIDS": [raster],
            "METHOD": 1,
            "NCLUSTER": n,
            "MAXITER": 0,
            "NORMALISE": False,
            "OLDVERSION": False,
            "UPDATEVIEW": False,
        },
        f"{layer_name}_kmeans_{n}",
        output_key="CLUSTER",
    )


def histogram_equalization(raster_layer):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(raster_layer)
    if layer is None:
        return {"error": f"Layer '{raster_layer}' not found"}
    try:
        from osgeo import gdal
        import numpy as np
    except ImportError as e:
        return {"error": f"GDAL/numpy not available: {e}"}

    try:
        src_path = layer.source()
        ds = gdal.Open(src_path)
        if ds is None:
            return {"error": f"Could not open raster: {src_path}"}

        out_path = _temp_raster_path()
        driver = gdal.GetDriverByName("GTiff")
        out_ds = driver.Create(
            out_path,
            ds.RasterXSize,
            ds.RasterYSize,
            ds.RasterCount,
            gdal.GDT_Byte,
        )
        out_ds.SetGeoTransform(ds.GetGeoTransform())
        out_ds.SetProjection(ds.GetProjection())

        for i in range(1, ds.RasterCount + 1):
            band = ds.GetRasterBand(i)
            arr = band.ReadAsArray().astype(np.float64)
            valid = ~np.isnan(arr)
            flat = arr[valid].astype(np.int64)
            if flat.size == 0:
                eq = np.zeros_like(arr, dtype=np.uint8)
            else:
                hist, bins = np.histogram(flat, bins=256, range=(flat.min(), flat.max() + 1))
                cdf = hist.cumsum()
                cdf_norm = (cdf - cdf.min()) * 255 / max(cdf.max() - cdf.min(), 1)
                cdf_norm = cdf_norm.astype(np.uint8)
                idx = np.clip(
                    ((arr - flat.min()) / max(flat.max() - flat.min(), 1) * 255).astype(np.int64),
                    0, 255,
                )
                eq = cdf_norm[idx]
            out_ds.GetRasterBand(i).WriteArray(eq)

        out_ds.FlushCache()
        out_ds = None
        ds = None

        new_name = f"{raster_layer}_equalized"
        new_layer = QgsRasterLayer(out_path, new_name)
        if not new_layer.isValid():
            return {"error": "Equalized raster is invalid"}
        QgsProject.instance().addMapLayer(new_layer)
        return {"success": True, "layer_name": new_name, "path": out_path}
    except Exception as e:
        return {"error": f"histogram_equalization failed: {e}"}


def mosaic_rasters(raster_layers_list):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    if not isinstance(raster_layers_list, list) or len(raster_layers_list) < 2:
        return {"error": "Provide a list of at least 2 raster layer names"}
    layers = []
    for name in raster_layers_list:
        layer = _find_layer_by_name(name)
        if layer is None:
            return {"error": f"Layer '{name}' not found"}
        layers.append(layer)
    return _run_raster_and_add(
        "gdal:merge",
        {
            "INPUT": layers,
            "PCT": False,
            "SEPARATE": False,
            "DATA_TYPE": 5,
        },
        "mosaic",
    )


def band_composite(layer_name, red_band, green_band, blue_band):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    raster = _find_layer_by_name(layer_name)
    if raster is None:
        return {"error": f"Layer '{layer_name}' not found"}
    try:
        r = int(red_band)
        g = int(green_band)
        b = int(blue_band)
    except (TypeError, ValueError):
        return {"error": "Band numbers must be integers"}
    return _run_raster_and_add(
        "gdal:translate",
        {
            "INPUT": raster,
            "TARGET_CRS": None,
            "NODATA": None,
            "COPY_SUBDATASETS": False,
            "OPTIONS": "",
            "EXTRA": f"-b {r} -b {g} -b {b}",
            "DATA_TYPE": 0,
        },
        f"{layer_name}_rgb_{r}{g}{b}",
    )


def _add_calculated_field(layer, field_name, expression_text):
    expr = QgsExpression(expression_text)
    if expr.hasParserError():
        return {"error": f"Invalid expression: {expr.parserErrorString()}"}

    provider = layer.dataProvider()
    if field_name not in [f.name() for f in layer.fields()]:
        provider.addAttributes([QgsField(field_name, QVariant.Double)])
        layer.updateFields()

    field_idx = layer.fields().indexFromName(field_name)
    context = QgsExpression.createContext() if hasattr(QgsExpression, "createContext") else None
    if context is None:
        from qgis.core import QgsExpressionContext, QgsExpressionContextUtils
        context = QgsExpressionContext()
        context.appendScopes(QgsExpressionContextUtils.globalProjectLayerScopes(layer))

    expr.prepare(context)

    layer.startEditing()
    try:
        for feature in layer.getFeatures():
            context.setFeature(feature)
            value = expr.evaluate(context)
            layer.changeAttributeValue(feature.id(), field_idx, float(value) if value is not None else None)
        layer.commitChanges()
    except Exception as e:
        layer.rollBack()
        return {"error": f"Field calculation failed: {e}"}

    return {"success": True, "layer_name": layer.name(), "field": field_name, "count": layer.featureCount()}


def calculate_area(layer_name):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    return _add_calculated_field(layer, "area", "$area")


def calculate_length(layer_name):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    return _add_calculated_field(layer, "length", "$length")


def centroid(layer_name):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    return _run_and_add(
        "native:centroids",
        {"INPUT": layer, "ALL_PARTS": False, "OUTPUT": "memory:"},
        f"{layer_name}_centroids",
    )


def reproject_layer(layer_name, crs_code):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    crs = QgsCoordinateReferenceSystem(crs_code)
    if not crs.isValid():
        return {"error": f"Invalid CRS: {crs_code}"}
    return _run_and_add(
        "native:reprojectlayer",
        {"INPUT": layer, "TARGET_CRS": crs, "OUTPUT": "memory:"},
        f"{layer_name}_{crs_code.replace(':', '_')}",
    )


def fix_geometries(layer_name):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    return _run_and_add(
        "native:fixgeometries",
        {"INPUT": layer, "OUTPUT": "memory:"},
        f"{layer_name}_fixed",
    )


def apply_categorized_style(layer_name, field):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    if field not in [f.name() for f in layer.fields()]:
        return {"error": f"Field '{field}' not found in '{layer_name}'"}

    unique_values = layer.uniqueValues(layer.fields().indexFromName(field))
    categories = []
    for value in unique_values:
        symbol = QgsSymbol.defaultSymbol(layer.geometryType())
        symbol.setColor(QColor(
            random.randint(40, 240),
            random.randint(40, 240),
            random.randint(40, 240),
        ))
        categories.append(QgsRendererCategory(value, symbol, str(value)))

    renderer = QgsCategorizedSymbolRenderer(field, categories)
    layer.setRenderer(renderer)
    layer.triggerRepaint()
    return {"success": True, "layer_name": layer_name, "categories": len(categories)}


def apply_graduated_style(layer_name, field):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    if field not in [f.name() for f in layer.fields()]:
        return {"error": f"Field '{field}' not found in '{layer_name}'"}

    try:
        renderer = QgsGraduatedSymbolRenderer(field, [])
        renderer.setClassificationMethod(QgsClassificationQuantile())
        renderer.updateClasses(layer, 5)
        layer.setRenderer(renderer)
        layer.triggerRepaint()
        return {"success": True, "layer_name": layer_name, "classes": 5}
    except Exception as e:
        return {"error": f"Graduated style failed: {e}"}


_FORMAT_DRIVERS = {
    "geojson": "GeoJSON",
    "shp": "ESRI Shapefile",
    "shapefile": "ESRI Shapefile",
    "esri shapefile": "ESRI Shapefile",
    "gpkg": "GPKG",
    "geopackage": "GPKG",
    "csv": "CSV",
    "kml": "KML",
    "gml": "GML",
}


def _write_vector(layer, output_path, driver_name, layer_options=None):
    options = QgsVectorFileWriter.SaveVectorOptions()
    options.driverName = driver_name
    if layer_options:
        options.layerOptions = layer_options
    transform_context = QgsProject.instance().transformContext()
    if hasattr(QgsVectorFileWriter, "writeAsVectorFormatV3"):
        result = QgsVectorFileWriter.writeAsVectorFormatV3(
            layer, output_path, transform_context, options
        )
    else:
        result = QgsVectorFileWriter.writeAsVectorFormatV2(
            layer, output_path, transform_context, options
        )
    code = result[0]
    if code != QgsVectorFileWriter.NoError:
        return {"error": f"Write failed (code {code}): {result[1] if len(result) > 1 else ''}"}
    return {"success": True, "path": output_path}


def export_layer(layer_name, output_path, format):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    driver = _FORMAT_DRIVERS.get(str(format).strip().lower())
    if driver is None:
        return {"error": f"Unsupported format: {format}. Supported: {sorted(set(_FORMAT_DRIVERS.values()))}"}
    return _write_vector(layer, output_path, driver)


def export_to_csv(layer_name, output_path):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    return _write_vector(
        layer,
        output_path,
        "CSV",
        layer_options=["GEOMETRY=AS_WKT", "SEPARATOR=COMMA"],
    )


def print_map(output_path):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    if iface is None:
        return {"error": "QGIS interface not available"}
    canvas = iface.mapCanvas()
    ext = output_path.lower().rsplit(".", 1)[-1] if "." in output_path else ""

    try:
        if ext in ("png", "jpg", "jpeg", "bmp"):
            canvas.saveAsImage(output_path)
            return {"success": True, "path": output_path, "format": ext}

        if ext == "pdf":
            from qgis.PyQt.QtPrintSupport import QPrinter
            from qgis.PyQt.QtGui import QPainter

            printer = QPrinter(QPrinter.HighResolution)
            printer.setOutputFormat(QPrinter.PdfFormat)
            printer.setOutputFileName(output_path)
            painter = QPainter(printer)
            try:
                canvas.render(painter)
            finally:
                painter.end()
            return {"success": True, "path": output_path, "format": "pdf"}

        return {"error": f"Unsupported output format: '.{ext}'. Use .png, .jpg, .bmp, or .pdf"}
    except Exception as e:
        return {"error": f"print_map failed: {e}"}


def select_by_attribute(layer_name, field, value):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    if field not in [f.name() for f in layer.fields()]:
        return {"error": f"Field '{field}' not found"}

    if isinstance(value, (int, float)):
        expr_text = f'"{field}" = {value}'
    else:
        escaped = str(value).replace("'", "''")
        expr_text = f"\"{field}\" = '{escaped}'"

    expr = QgsExpression(expr_text)
    if expr.hasParserError():
        return {"error": f"Invalid expression: {expr.parserErrorString()}"}
    layer.selectByExpression(expr_text)
    return {"success": True, "count": layer.selectedFeatureCount(), "expression": expr_text}


def export_attribute_table(layer_name, output_path):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    return _write_vector(
        layer,
        output_path,
        "CSV",
        layer_options=["GEOMETRY=AS_WKT", "SEPARATOR=COMMA"],
    )


def field_calculator(layer_name, new_field, expression):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    return _add_calculated_field(layer, new_field, expression)


def get_feature_count(layer_name):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    return {"layer_name": layer_name, "count": layer.featureCount()}


def filter_features(layer_name, expression):
    return run_query(layer_name, expression)


def pan_sharpening(ms_layer, pan_layer):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    ms = _find_layer_by_name(ms_layer)
    pan = _find_layer_by_name(pan_layer)
    if ms is None:
        return {"error": f"Layer '{ms_layer}' not found"}
    if pan is None:
        return {"error": f"Layer '{pan_layer}' not found"}
    return _run_raster_and_add(
        "gdal:pansharp",
        {
            "SPECTRAL": ms,
            "PANCHROMATIC": pan,
            "RESAMPLING": 2,
        },
        f"{ms_layer}_pansharp",
    )


def supervised_classification(layer_name, training_layer):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    raster = _find_layer_by_name(layer_name)
    training = _find_layer_by_name(training_layer)
    if raster is None:
        return {"error": f"Layer '{layer_name}' not found"}
    if training is None:
        return {"error": f"Layer '{training_layer}' not found"}
    return _run_raster_and_add(
        "saga:supervisedclassificationforgrids",
        {
            "GRIDS": [raster],
            "TRAINING": training,
            "TRAINING_CLASS": "class",
            "METHOD": 2,
            "NORMALISE": False,
        },
        f"{layer_name}_classified",
        output_key="CLASSES",
    )


def calculate_ndre(red_edge_layer, nir_layer):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    re = _find_layer_by_name(red_edge_layer)
    nir = _find_layer_by_name(nir_layer)
    if re is None:
        return {"error": f"Layer '{red_edge_layer}' not found"}
    if nir is None:
        return {"error": f"Layer '{nir_layer}' not found"}
    return _run_raster_and_add(
        "gdal:rastercalculator",
        {
            "INPUT_A": nir,
            "BAND_A": 1,
            "INPUT_B": re,
            "BAND_B": 1,
            "FORMULA": "(A.astype(float)-B)/(A.astype(float)+B)",
            "NO_DATA": None,
            "RTYPE": 5,
        },
        "NDRE",
    )


def apply_heatmap_style(layer_name, field):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    try:
        renderer = QgsHeatmapRenderer()
        if field and field in [f.name() for f in layer.fields()]:
            renderer.setWeightExpression(f'"{field}"')
        renderer.setRadius(10)
        layer.setRenderer(renderer)
        layer.triggerRepaint()
        return {"success": True, "layer_name": layer_name}
    except Exception as e:
        return {"error": f"Heatmap style failed: {e}"}


def change_layer_color(layer_name, color_hex):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    color = QColor(color_hex)
    if not color.isValid():
        return {"error": f"Invalid color: {color_hex}"}
    try:
        renderer = layer.renderer()
        if renderer is None or not hasattr(renderer, "symbol"):
            symbol = QgsSymbol.defaultSymbol(layer.geometryType())
            symbol.setColor(color)
            layer.setRenderer(QgsSingleSymbolRenderer(symbol))
        else:
            renderer.symbol().setColor(color)
        layer.triggerRepaint()
        return {"success": True, "layer_name": layer_name, "color": color_hex}
    except Exception as e:
        return {"error": f"change_layer_color failed: {e}"}


def open_attribute_table(layer_name):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    if iface is None:
        return {"error": "QGIS interface not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    iface.showAttributeTable(layer)
    return {"success": True, "layer_name": layer_name}


def zoom_to_feature(layer_name, feature_id):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    if iface is None:
        return {"error": "QGIS interface not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    try:
        fid = int(feature_id)
    except (TypeError, ValueError):
        return {"error": f"feature_id must be an integer, got {feature_id!r}"}
    feature = layer.getFeature(fid)
    if not feature.isValid():
        return {"error": f"Feature id {fid} not found in '{layer_name}'"}
    geom = feature.geometry()
    if geom is None or geom.isEmpty():
        return {"error": "Feature has no geometry"}
    canvas = iface.mapCanvas()
    canvas.setExtent(geom.boundingBox())
    canvas.refresh()
    layer.selectByIds([fid])
    return {"success": True, "layer_name": layer_name, "feature_id": fid}


def generate_report(title, content):
    import os
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(str(title), 0)
        for paragraph in str(content).split("\n\n"):
            doc.add_paragraph(paragraph)
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        if not os.path.isdir(desktop):
            desktop = os.path.expanduser("~")
        safe_title = "".join(c if c.isalnum() or c in "._- " else "_" for c in str(title)).strip() or "report"
        out_path = os.path.join(desktop, f"{safe_title}.docx")
        doc.save(out_path)
        return {"success": True, "path": out_path}
    except ImportError:
        return {
            "error": (
                "python-docx not installed. Run in OSGeo4W Shell: "
                "python -m pip install python-docx"
            )
        }
    except Exception as e:
        return {"error": f"Error: {e}"}


def generate_csv(layer_name, output_path):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    return _write_vector(
        layer,
        output_path,
        "CSV",
        layer_options=["GEOMETRY=AS_WKT", "SEPARATOR=COMMA"],
    )


def generate_map_image(output_path):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    if iface is None:
        return {"error": "QGIS interface not available"}
    try:
        iface.mapCanvas().saveAsImage(output_path)
        return {"success": True, "path": output_path}
    except Exception as e:
        return {"error": f"generate_map_image failed: {e}"}


def get_crs(layer_name):
    if not QGIS_AVAILABLE:
        return {"error": "QGIS not available"}
    layer = _find_layer_by_name(layer_name)
    if layer is None:
        return {"error": f"Layer '{layer_name}' not found"}
    crs = layer.crs()
    return {
        "layer_name": layer_name,
        "authid": crs.authid(),
        "description": crs.description(),
        "is_geographic": crs.isGeographic(),
        "units": crs.mapUnits(),
    }


TOOLS_SCHEMA = [
    {
        "type": "function",
        "function": {
            "name": "get_layers",
            "description": "Get all layers in the current QGIS project. Returns a list of layers with name, type, and id.",
            "parameters": {
                "type": "object",
                "properties": {},
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_attributes",
            "description": "Get the list of field/attribute names for a given layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {
                        "type": "string",
                        "description": "Name of the layer to inspect.",
                    },
                },
                "required": ["layer_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "run_query",
            "description": "Filter features of a layer using a QGIS expression (subset string).",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {
                        "type": "string",
                        "description": "Name of the layer to filter.",
                    },
                    "expression": {
                        "type": "string",
                        "description": "QGIS expression, e.g. \"population > 1000\".",
                    },
                },
                "required": ["layer_name", "expression"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "buffer_analysis",
            "description": "Create a buffer around all features of a layer at a given distance (in layer CRS units). Adds the buffer as a new layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {
                        "type": "string",
                        "description": "Name of the input layer.",
                    },
                    "distance": {
                        "type": "number",
                        "description": "Buffer distance in the layer's CRS units.",
                    },
                },
                "required": ["layer_name", "distance"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "highlight_features",
            "description": "Select (highlight) features in a layer that match a QGIS expression. Returns count of selected features.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {
                        "type": "string",
                        "description": "Name of the layer.",
                    },
                    "expression": {
                        "type": "string",
                        "description": "QGIS expression to select features by.",
                    },
                },
                "required": ["layer_name", "expression"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "remove_layer",
            "description": "Remove a layer from the current QGIS project by name.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer to remove."},
                },
                "required": ["layer_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "rename_layer",
            "description": "Rename a layer in the current QGIS project.",
            "parameters": {
                "type": "object",
                "properties": {
                    "old_name": {"type": "string", "description": "Current layer name."},
                    "new_name": {"type": "string", "description": "New layer name."},
                },
                "required": ["old_name", "new_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "zoom_to_layer",
            "description": "Zoom the QGIS map canvas to the extent of a layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer to zoom to."},
                },
                "required": ["layer_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "toggle_visibility",
            "description": "Show or hide a layer in the QGIS layer tree.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer."},
                    "visible": {"type": "boolean", "description": "True to show, False to hide. Defaults to True."},
                },
                "required": ["layer_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "add_layer_from_path",
            "description": "Load a vector or raster layer from a file path and add it to the current QGIS project.",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Absolute path to a vector (e.g. .shp, .geojson, .gpkg) or raster (e.g. .tif) file."},
                },
                "required": ["file_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "clip_layer",
            "description": "Clip a vector layer by a mask (overlay) layer. Adds the clipped result as a new layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "input_layer": {"type": "string", "description": "Name of the layer to clip."},
                    "mask_layer": {"type": "string", "description": "Name of the mask/overlay layer."},
                },
                "required": ["input_layer", "mask_layer"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "intersect_layers",
            "description": "Compute the geometric intersection of two vector layers. Adds the result as a new layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer1": {"type": "string", "description": "Name of the first layer."},
                    "layer2": {"type": "string", "description": "Name of the second layer."},
                },
                "required": ["layer1", "layer2"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "union_layers",
            "description": "Compute the geometric union of two vector layers. Adds the result as a new layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer1": {"type": "string", "description": "Name of the first layer."},
                    "layer2": {"type": "string", "description": "Name of the second layer."},
                },
                "required": ["layer1", "layer2"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "dissolve_layer",
            "description": "Dissolve features of a vector layer, optionally grouped by a field. If no field is given, all features are dissolved together.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer to dissolve."},
                    "field": {"type": "string", "description": "Optional field name to dissolve by."},
                },
                "required": ["layer_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "merge_layers",
            "description": "Merge multiple vector layers into a single layer. Requires at least 2 layer names.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_names_list": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "List of layer names to merge.",
                    },
                },
                "required": ["layer_names_list"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "spatial_join",
            "description": "Join attributes from one layer to another by spatial location (intersects). Adds the result as a new layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "target_layer": {"type": "string", "description": "Layer that receives the joined attributes."},
                    "join_layer": {"type": "string", "description": "Layer whose attributes will be joined."},
                },
                "required": ["target_layer", "join_layer"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "calculate_ndvi",
            "description": "Compute NDVI = (NIR - Red) / (NIR + Red) from two raster bands. Adds the result as a new raster layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "red_layer": {"type": "string", "description": "Name of the Red band raster layer."},
                    "nir_layer": {"type": "string", "description": "Name of the NIR band raster layer."},
                },
                "required": ["red_layer", "nir_layer"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "calculate_ndwi",
            "description": "Compute NDWI = (Green - NIR) / (Green + NIR) from two raster bands. Adds the result as a new raster layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "green_layer": {"type": "string", "description": "Name of the Green band raster layer."},
                    "nir_layer": {"type": "string", "description": "Name of the NIR band raster layer."},
                },
                "required": ["green_layer", "nir_layer"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "hillshade",
            "description": "Compute a hillshade from a DEM raster. Adds the result as a new raster layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "dem_layer": {"type": "string", "description": "Name of the DEM raster layer."},
                },
                "required": ["dem_layer"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "slope_analysis",
            "description": "Compute slope from a DEM raster. Adds the result as a new raster layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "dem_layer": {"type": "string", "description": "Name of the DEM raster layer."},
                },
                "required": ["dem_layer"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "aspect_analysis",
            "description": "Compute aspect from a DEM raster. Adds the result as a new raster layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "dem_layer": {"type": "string", "description": "Name of the DEM raster layer."},
                },
                "required": ["dem_layer"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "zonal_statistics",
            "description": "Compute zonal statistics (count, sum, mean, median, stdev, min, max, range) of a raster within polygons of a vector layer. Adds a new vector layer with the statistics as attributes.",
            "parameters": {
                "type": "object",
                "properties": {
                    "raster_layer": {"type": "string", "description": "Name of the raster layer to summarise."},
                    "vector_layer": {"type": "string", "description": "Name of the polygon vector layer defining zones."},
                },
                "required": ["raster_layer", "vector_layer"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "raster_clip",
            "description": "Clip a raster layer by a vector mask layer. Adds the clipped raster as a new layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "raster_layer": {"type": "string", "description": "Name of the raster layer to clip."},
                    "mask_layer": {"type": "string", "description": "Name of the vector mask layer."},
                },
                "required": ["raster_layer", "mask_layer"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "unsupervised_classification",
            "description": "Perform K-means unsupervised classification on a raster layer (uses SAGA). Adds the result as a new raster layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the raster layer to classify."},
                    "num_classes": {"type": "integer", "description": "Number of clusters/classes."},
                },
                "required": ["layer_name", "num_classes"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "histogram_equalization",
            "description": "Apply histogram equalization to a raster layer (per band). Adds the equalized raster as a new layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "raster_layer": {"type": "string", "description": "Name of the raster layer."},
                },
                "required": ["raster_layer"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "mosaic_rasters",
            "description": "Mosaic multiple raster layers into a single raster. Requires at least 2 layers.",
            "parameters": {
                "type": "object",
                "properties": {
                    "raster_layers_list": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "List of raster layer names to mosaic.",
                    },
                },
                "required": ["raster_layers_list"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "band_composite",
            "description": "Create an RGB composite by selecting three bands from a multi-band raster layer. Adds the result as a new raster layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the multi-band raster layer."},
                    "red_band": {"type": "integer", "description": "Band number for the Red channel (1-based)."},
                    "green_band": {"type": "integer", "description": "Band number for the Green channel (1-based)."},
                    "blue_band": {"type": "integer", "description": "Band number for the Blue channel (1-based)."},
                },
                "required": ["layer_name", "red_band", "green_band", "blue_band"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "calculate_area",
            "description": "Calculate the area of every feature in a polygon layer and write it to an 'area' field (created if missing).",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the polygon layer."},
                },
                "required": ["layer_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "calculate_length",
            "description": "Calculate the length of every feature in a line layer and write it to a 'length' field (created if missing).",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the line layer."},
                },
                "required": ["layer_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "centroid",
            "description": "Create a centroid point layer from a polygon (or line) layer. Adds the result as a new layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the input layer."},
                },
                "required": ["layer_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "reproject_layer",
            "description": "Reproject a vector layer to a target CRS (e.g. 'EPSG:4326'). Adds the reprojected layer as a new layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer to reproject."},
                    "crs_code": {"type": "string", "description": "Target CRS, e.g. 'EPSG:4326'."},
                },
                "required": ["layer_name", "crs_code"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fix_geometries",
            "description": "Fix invalid geometries in a vector layer. Adds the cleaned layer as a new layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the input layer."},
                },
                "required": ["layer_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "apply_categorized_style",
            "description": "Apply a categorized renderer to a vector layer, one category per unique value of the given field.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the vector layer."},
                    "field": {"type": "string", "description": "Field name to categorize by."},
                },
                "required": ["layer_name", "field"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "apply_graduated_style",
            "description": "Apply a graduated (quantile, 5 classes) renderer to a vector layer based on a numeric field.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the vector layer."},
                    "field": {"type": "string", "description": "Numeric field to classify by."},
                },
                "required": ["layer_name", "field"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "export_layer",
            "description": "Export a vector layer to a file. Supported formats: 'geojson', 'shp', 'gpkg', 'kml', 'gml'.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer to export."},
                    "output_path": {"type": "string", "description": "Absolute output file path."},
                    "format": {"type": "string", "description": "File format: geojson, shp, gpkg, kml, or gml."},
                },
                "required": ["layer_name", "output_path", "format"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "export_to_csv",
            "description": "Export the attribute table of a layer to CSV (geometry written as WKT).",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer."},
                    "output_path": {"type": "string", "description": "Absolute output .csv path."},
                },
                "required": ["layer_name", "output_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "print_map",
            "description": "Export the current QGIS map canvas to an image (.png, .jpg, .bmp) or PDF.",
            "parameters": {
                "type": "object",
                "properties": {
                    "output_path": {"type": "string", "description": "Absolute output path ending in .png, .jpg, .bmp, or .pdf."},
                },
                "required": ["output_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "select_by_attribute",
            "description": "Select features in a layer where a field equals a value.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer."},
                    "field": {"type": "string", "description": "Field name."},
                    "value": {"description": "Value to match (string or number)."},
                },
                "required": ["layer_name", "field", "value"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "export_attribute_table",
            "description": "Export the full attribute table of a layer to a CSV file (geometry as WKT).",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer."},
                    "output_path": {"type": "string", "description": "Absolute output .csv path."},
                },
                "required": ["layer_name", "output_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "field_calculator",
            "description": "Add a new field to a vector layer and populate it using a QGIS expression.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer."},
                    "new_field": {"type": "string", "description": "Name of the new field."},
                    "expression": {"type": "string", "description": "QGIS expression to evaluate per feature."},
                },
                "required": ["layer_name", "new_field", "expression"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_feature_count",
            "description": "Return the number of features in a layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer."},
                },
                "required": ["layer_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "filter_features",
            "description": "Filter a vector layer by a QGIS expression (applies a subset string).",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer."},
                    "expression": {"type": "string", "description": "QGIS expression."},
                },
                "required": ["layer_name", "expression"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "pan_sharpening",
            "description": "Pan-sharpen a multispectral raster using a higher-resolution panchromatic raster (gdal:pansharp).",
            "parameters": {
                "type": "object",
                "properties": {
                    "ms_layer": {"type": "string", "description": "Multispectral raster layer."},
                    "pan_layer": {"type": "string", "description": "Panchromatic raster layer."},
                },
                "required": ["ms_layer", "pan_layer"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "supervised_classification",
            "description": "Supervised classification of a raster using a training vector layer with a 'class' field (uses SAGA).",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Raster layer to classify."},
                    "training_layer": {"type": "string", "description": "Training polygon layer with a 'class' attribute."},
                },
                "required": ["layer_name", "training_layer"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "calculate_ndre",
            "description": "Compute NDRE = (NIR - RedEdge) / (NIR + RedEdge) from two raster bands. Adds the result as a new raster layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "red_edge_layer": {"type": "string", "description": "Red Edge band raster layer."},
                    "nir_layer": {"type": "string", "description": "NIR band raster layer."},
                },
                "required": ["red_edge_layer", "nir_layer"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "apply_heatmap_style",
            "description": "Apply a heatmap renderer to a point vector layer, optionally weighted by a field.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Point vector layer."},
                    "field": {"type": "string", "description": "Optional weight field name."},
                },
                "required": ["layer_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "change_layer_color",
            "description": "Change the fill/line color of a vector layer's symbol to a hex color (e.g. '#ff0000').",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer."},
                    "color_hex": {"type": "string", "description": "Hex color string like '#ff0000'."},
                },
                "required": ["layer_name", "color_hex"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "open_attribute_table",
            "description": "Open the attribute table dialog for a layer in the QGIS GUI.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer."},
                },
                "required": ["layer_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "zoom_to_feature",
            "description": "Zoom the map canvas to a specific feature by its feature id and select it.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer."},
                    "feature_id": {"type": "integer", "description": "QGIS feature id (FID)."},
                },
                "required": ["layer_name", "feature_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_report",
            "description": "Generate a Word (.docx) report with a title and body content. The file is saved to the user's Desktop.",
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "Report title (also used as filename)."},
                    "content": {"type": "string", "description": "Body text. Use blank lines to separate paragraphs."},
                },
                "required": ["title", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_csv",
            "description": "Export a layer's attribute table to a CSV file (geometry written as WKT).",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer."},
                    "output_path": {"type": "string", "description": "Absolute output .csv path."},
                },
                "required": ["layer_name", "output_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_map_image",
            "description": "Export the current QGIS map canvas as a PNG image.",
            "parameters": {
                "type": "object",
                "properties": {
                    "output_path": {"type": "string", "description": "Absolute output .png path."},
                },
                "required": ["output_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_crs",
            "description": "Return the CRS authid, description, geographic flag, and map units for a layer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "layer_name": {"type": "string", "description": "Name of the layer."},
                },
                "required": ["layer_name"],
            },
        },
    },
]

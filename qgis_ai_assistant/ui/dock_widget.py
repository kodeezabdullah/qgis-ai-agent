import threading
import traceback

from qgis.PyQt.QtCore import Qt, QObject, QUrl, pyqtSignal, pyqtSlot
from qgis.PyQt.QtWidgets import (
    QDockWidget,
    QWidget,
    QVBoxLayout,
    QLabel,
    QInputDialog,
    QLineEdit,
    QFileDialog,
)
from qgis.core import QgsSettings


SETTINGS_KEY = "qgis_ai_agent/api_key"
ONBOARDING_KEY = "qgis_ai_agent/onboarding_complete"


def _dock_areas_left_right():
    left = getattr(Qt, "LeftDockWidgetArea", None)
    right = getattr(Qt, "RightDockWidgetArea", None)
    if left is None or right is None:
        left = Qt.DockWidgetArea.LeftDockWidgetArea
        right = Qt.DockWidgetArea.RightDockWidgetArea
    return left | right


def _align_center():
    return getattr(Qt, "AlignCenter", None) or Qt.AlignmentFlag.AlignCenter


WEBENGINE_AVAILABLE = False
QWebEngineView = None
QWebChannel = None
_WEBENGINE_IMPORT_ERROR = None
try:
    from qgis.PyQt.QtWebEngineWidgets import QWebEngineView  # noqa: F401
    from qgis.PyQt.QtWebChannel import QWebChannel  # noqa: F401
    WEBENGINE_AVAILABLE = True
except Exception as e:
    _WEBENGINE_IMPORT_ERROR = e
    print(f"[QgisAiAgentDock] QtWebEngine/QtWebChannel import failed: {e}")


CHAT_HTML = r"""
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8" />
<title>QGIS AI Agent</title>
<script src="qrc:///qtwebchannel/qwebchannel.js"></script>
<script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
<style>
  :root {
    --bg: #f7f8fa;
    --header-from: #e3f0ff;
    --header-to: #d4e7ff;
    --header-text: #0b3d77;
    --user-bg: #1a73e8;
    --user-text: #ffffff;
    --ai-bg: #ffffff;
    --ai-border: #e3e4e8;
    --ai-text: #1f2024;
    --muted: #6a6f76;
    --accent: #1a73e8;
    --accent-hover: #1666d2;
    --input-bg: #ffffff;
    --input-border: #d8d9dd;
    --code-bg: #f1f3f6;
  }
  * { box-sizing: border-box; }
  html, body {
    margin: 0; padding: 0; height: 100%;
    background: var(--bg); color: var(--ai-text);
    font-family: "Inter", -apple-system, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
    font-size: 14px;
  }
  body { display: flex; flex-direction: column; }

  header {
    background: linear-gradient(135deg, var(--header-from), var(--header-to));
    color: var(--header-text);
    padding: 14px 16px;
    font-weight: 600;
    font-size: 15px;
    display: flex; align-items: center; justify-content: space-between;
    border-bottom: 1px solid #cfdbeb;
  }
  header .title { display: flex; align-items: center; gap: 8px; }
  header .gear {
    cursor: pointer; opacity: 0.7; font-size: 16px;
    background: transparent; border: 0; color: var(--header-text);
  }
  header .gear:hover { opacity: 1; }

  #chat {
    flex: 1; overflow-y: auto;
    padding: 16px;
    display: flex; flex-direction: column; gap: 12px;
    scroll-behavior: smooth;
  }
  .bubble {
    max-width: 88%;
    padding: 10px 14px;
    border-radius: 14px;
    line-height: 1.45;
    word-wrap: break-word;
    white-space: pre-wrap;
  }
  .bubble.user {
    align-self: flex-end;
    background: var(--user-bg);
    color: var(--user-text);
    border-bottom-right-radius: 4px;
  }
  .bubble.ai {
    align-self: flex-start;
    background: var(--ai-bg);
    color: var(--ai-text);
    border: 1px solid var(--ai-border);
    border-bottom-left-radius: 4px;
    box-shadow: 0 1px 2px rgba(0,0,0,0.03);
  }
  .bubble.ai p { margin: 0 0 8px; }
  .bubble.ai p:last-child { margin-bottom: 0; }
  .bubble.ai code {
    background: var(--code-bg);
    padding: 1px 5px; border-radius: 4px;
    font-family: "JetBrains Mono", Consolas, monospace;
    font-size: 12.5px;
  }
  .bubble.ai pre {
    background: var(--code-bg);
    padding: 10px 12px; border-radius: 8px;
    overflow-x: auto; margin: 6px 0;
  }
  .bubble.ai pre code { background: transparent; padding: 0; }
  .bubble.ai table {
    border-collapse: collapse; margin: 6px 0;
  }
  .bubble.ai th, .bubble.ai td {
    border: 1px solid var(--ai-border);
    padding: 4px 8px;
  }

  .typing {
    align-self: flex-start;
    display: inline-flex; align-items: center; gap: 4px;
    padding: 10px 14px;
    background: var(--ai-bg);
    border: 1px solid var(--ai-border);
    border-radius: 14px;
    border-bottom-left-radius: 4px;
  }
  .typing span {
    width: 6px; height: 6px; border-radius: 50%;
    background: #9aa0a6;
    animation: bounce 1.2s infinite ease-in-out both;
  }
  .typing span:nth-child(2) { animation-delay: 0.15s; }
  .typing span:nth-child(3) { animation-delay: 0.3s; }
  @keyframes bounce {
    0%, 80%, 100% { transform: scale(0.6); opacity: 0.4; }
    40% { transform: scale(1); opacity: 1; }
  }

  #status {
    min-height: 18px;
    padding: 0 16px;
    font-size: 12px;
    color: var(--muted);
  }

  #input-row {
    display: flex; align-items: flex-end; gap: 8px;
    padding: 10px 12px;
    background: var(--bg);
    border-top: 1px solid var(--ai-border);
  }
  #attach, #send {
    background: transparent;
    border: none;
    cursor: pointer;
    width: 36px; height: 36px;
    border-radius: 50%;
    display: flex; align-items: center; justify-content: center;
    font-size: 18px;
    color: var(--muted);
  }
  #attach:hover { background: #eef0f3; color: var(--accent); }
  #send {
    background: var(--accent); color: white;
  }
  #send:hover { background: var(--accent-hover); }
  #send:disabled { background: #aac5ee; cursor: not-allowed; }

  #composer {
    flex: 1;
    display: flex; align-items: center;
    background: var(--input-bg);
    border: 1px solid var(--input-border);
    border-radius: 18px;
    padding: 4px 10px;
  }
  #input {
    flex: 1;
    resize: none;
    border: 0;
    outline: 0;
    background: transparent;
    font-family: inherit;
    font-size: 14px;
    line-height: 1.4;
    max-height: 140px;
    padding: 6px 4px;
    min-height: 22px;
  }

  footer {
    text-align: center;
    padding: 8px 12px 10px;
    font-size: 11px;
    color: var(--muted);
    background: var(--bg);
    border-top: 1px solid var(--ai-border);
  }

  .wizard {
    align-self: stretch;
    background: var(--ai-bg);
    border: 1px solid var(--ai-border);
    border-radius: 14px;
    padding: 14px 16px;
    box-shadow: 0 1px 2px rgba(0,0,0,0.03);
  }
  .wizard h3 { margin: 0 0 8px; font-size: 15px; }
  .wizard p { margin: 0 0 8px; line-height: 1.5; white-space: pre-wrap; }
  .wizard .btn-row {
    display: flex; gap: 8px; flex-wrap: wrap; margin-top: 10px;
  }
  .wizard button {
    border: 1px solid var(--accent);
    background: var(--accent);
    color: white;
    padding: 6px 14px;
    border-radius: 18px;
    font-size: 13px;
    cursor: pointer;
  }
  .wizard button.secondary {
    background: white;
    color: var(--accent);
  }
  .wizard button:hover { background: var(--accent-hover); color: white; }
  .wizard .step-meta {
    font-size: 11px; color: var(--muted); margin-bottom: 6px;
  }
</style>
</head>
<body>
  <header>
    <div class="title">🗺️ QGIS AI Agent</div>
    <button class="gear" id="settings" title="Settings">⚙</button>
  </header>

  <div id="chat"></div>
  <div id="status"></div>

  <div id="input-row">
    <button id="attach" title="Attach file">📎</button>
    <div id="composer">
      <textarea id="input" rows="1" placeholder="Ask me anything about your layers..."></textarea>
    </div>
    <button id="send" title="Send">➤</button>
  </div>

  <footer>Built by kodeezabdullah | Geoinformatics Engineering — NUST IGIS</footer>

<script>
  const chat = document.getElementById("chat");
  const input = document.getElementById("input");
  const sendBtn = document.getElementById("send");
  const attachBtn = document.getElementById("attach");
  const settingsBtn = document.getElementById("settings");
  const statusEl = document.getElementById("status");

  let bridge = null;
  let typingEl = null;

  function scrollToBottom() {
    chat.scrollTop = chat.scrollHeight;
  }

  function renderMarkdown(text) {
    if (window.marked) {
      return marked.parse(text);
    }
    const esc = text
      .replace(/&/g, "&amp;").replace(/</g, "&lt;").replace(/>/g, "&gt;");
    return esc.replace(/\n/g, "<br>");
  }

  function addBubble(role, text) {
    const div = document.createElement("div");
    div.className = "bubble " + role;
    if (role === "ai") {
      div.innerHTML = renderMarkdown(text);
    } else {
      div.textContent = text;
    }
    chat.appendChild(div);
    scrollToBottom();
  }

  function showTyping() {
    if (typingEl) return;
    typingEl = document.createElement("div");
    typingEl.className = "typing";
    typingEl.innerHTML = "<span></span><span></span><span></span>";
    chat.appendChild(typingEl);
    scrollToBottom();
  }

  function hideTyping() {
    if (typingEl && typingEl.parentNode) {
      typingEl.parentNode.removeChild(typingEl);
    }
    typingEl = null;
  }

  function setStatus(text) {
    statusEl.textContent = text || "";
  }

  function autoResize() {
    input.style.height = "auto";
    input.style.height = Math.min(input.scrollHeight, 140) + "px";
  }
  input.addEventListener("input", autoResize);

  function sendMessage() {
    const text = input.value.trim();
    if (!text) return;
    addBubble("user", text);
    input.value = "";
    autoResize();
    showTyping();
    setStatus("Thinking...");
    sendBtn.disabled = true;
    if (bridge) {
      bridge.sendMessage(text);
    } else {
      hideTyping();
      setStatus("");
      addBubble("ai", "_Bridge not ready yet — please retry in a moment._");
      sendBtn.disabled = false;
    }
  }

  sendBtn.addEventListener("click", sendMessage);
  input.addEventListener("keydown", (e) => {
    if (e.key === "Enter" && !e.shiftKey) {
      e.preventDefault();
      sendMessage();
    }
  });
  attachBtn.addEventListener("click", () => {
    if (bridge && bridge.attachFile) bridge.attachFile();
  });
  settingsBtn.addEventListener("click", () => {
    if (bridge && bridge.openSettings) bridge.openSettings();
  });

  function receiveMessage(text) {
    hideTyping();
    setStatus("");
    sendBtn.disabled = false;
    addBubble("ai", text);
  }

  function welcome() {
    addBubble(
      "ai",
      "Hello! I'm your QGIS AI Assistant 🗺️\n\n" +
      "I can help you analyze spatial data, run queries, create buffers, calculate NDVI, and much more!\n\n" +
      "Type your question or command below."
    );
  }

  function addFileChip(filename) {
    const div = document.createElement("div");
    div.className = "bubble ai";
    div.innerHTML =
      "📎 <b>File attached:</b> " + filename.replace(/</g, "&lt;") +
      "<br><span style='opacity:0.7;font-size:12px;'>Analyzing...</span>";
    chat.appendChild(div);
    showTyping();
    scrollToBottom();
  }

  let wizardEl = null;
  function clearWizard() {
    if (wizardEl && wizardEl.parentNode) wizardEl.parentNode.removeChild(wizardEl);
    wizardEl = null;
  }

  function renderWizard(step) {
    clearWizard();
    const card = document.createElement("div");
    card.className = "wizard";

    const steps = {
      1: {
        meta: "Welcome",
        html: "<h3>👋 Welcome to QGIS AI Agent!</h3>" +
              "<p>Before we start, you need a FREE API key from OpenRouter.\n" +
              "This takes less than 2 minutes! Let's go step by step 🚀</p>",
        buttons: [{ label: "Get Started", action: () => renderWizard(2) }],
      },
      2: {
        meta: "Step 1 of 4",
        html: "<h3>📋 Step 1: Create your FREE OpenRouter account</h3>" +
              "<p>1. Click 'Open OpenRouter' below\n" +
              "2. Click 'Sign Up' (top right)\n" +
              "3. Enter your email and create password\n" +
              "4. Verify your email\n\n" +
              "✅ Done? Click Next!</p>",
        buttons: [
          { label: "Open OpenRouter", action: () => bridge.openUrl("https://openrouter.ai/"), secondary: true },
          { label: "Next →", action: () => renderWizard(3) },
        ],
      },
      3: {
        meta: "Step 2 of 4",
        html: "<h3>🔑 Step 2: Create your API Key</h3>" +
              "<p>1. After login, click your profile (top right)\n" +
              "2. Go to 'API Keys'\n" +
              "3. Click '+ New Key'\n" +
              "4. Name it: 'qgis-ai-agent'\n" +
              "5. Click 'Create'\n" +
              "6. COPY the key (you only see it once!)</p>",
        buttons: [
          { label: "Open API Keys Page", action: () => bridge.openUrl("https://openrouter.ai/keys"), secondary: true },
          { label: "Next →", action: () => renderWizard(4) },
        ],
      },
      4: {
        meta: "Step 3 of 4",
        html: "<h3>⚙️ Step 3: Paste your API Key</h3>" +
              "<p>Click the ⚙️ Settings button (top right of this panel).\n" +
              "Paste your API key and click OK.</p>",
        buttons: [
          { label: "Open Settings ⚙️", action: () => bridge.openSettings(), secondary: true },
          { label: "Next →", action: () => renderWizard(5) },
        ],
      },
      5: {
        meta: "All done",
        html: "<h3>🎉 You're all set!</h3>" +
              "<p>Your QGIS AI Assistant is ready to use!\n\n" +
              "Try asking:\n" +
              "- 'List all my layers'\n" +
              "- 'Create a buffer around my layer'\n" +
              "- 'Calculate area of my features'\n" +
              "- 'Filter features by attribute'\n\n" +
              "What would you like to do? 🗺️</p>",
        buttons: [
          { label: "Start chatting", action: () => {
              bridge.markOnboardingComplete();
              clearWizard();
              input.focus();
            }
          },
        ],
      },
    };

    const cfg = steps[step] || steps[1];
    card.innerHTML =
      `<div class="step-meta">${cfg.meta}</div>` + cfg.html +
      `<div class="btn-row" id="btn-row"></div>`;
    chat.appendChild(card);
    const row = card.querySelector("#btn-row");
    cfg.buttons.forEach(b => {
      const btn = document.createElement("button");
      btn.textContent = b.label;
      if (b.secondary) btn.classList.add("secondary");
      btn.addEventListener("click", b.action);
      row.appendChild(btn);
    });
    wizardEl = card;
    scrollToBottom();
  }

  new QWebChannel(qt.webChannelTransport, function(channel) {
    bridge = channel.objects.bridge;
    bridge.receiveMessage.connect(receiveMessage);
    bridge.setStatus.connect(setStatus);
    if (bridge.fileAttached) bridge.fileAttached.connect(addFileChip);
    if (bridge.apiKeySaved) bridge.apiKeySaved.connect(() => {
      if (wizardEl) renderWizard(5);
    });
    bridge.getInitialState(function(state) {
      if (state === "onboarding") {
        renderWizard(1);
      } else {
        welcome();
      }
    });
  });
</script>
</body>
</html>
"""


def _read_attached_file(path):
    import os
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return {"text": text, "is_image": False}, None

        if ext == ".docx":
            from docx import Document
            doc = Document(path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            return {"text": text, "is_image": False}, None

        if ext in (".png", ".jpg", ".jpeg", ".bmp", ".gif"):
            import base64
            with open(path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode()
            mime = "jpeg" if ext in (".jpg", ".jpeg") else ext.lstrip(".")
            return {"text": None, "is_image": True, "b64": b64, "mime": mime}, None

        if ext == ".csv":
            import pandas as pd
            df = pd.read_csv(path)
            text = (
                f"CSV with {len(df)} rows, {len(df.columns)} columns\n"
                f"Columns: {list(df.columns)}\n"
                f"First 5 rows:\n{df.head().to_string()}"
            )
            return {"text": text, "is_image": False}, None

        if ext in (".xlsx", ".xls"):
            import pandas as pd
            df = pd.read_excel(path)
            text = (
                f"Excel with {len(df)} rows, {len(df.columns)} columns\n"
                f"Columns: {list(df.columns)}\n"
                f"First 5 rows:\n{df.head().to_string()}"
            )
            return {"text": text, "is_image": False}, None

        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return {"text": f.read(), "is_image": False}, None
    except ImportError as e:
        return None, (
            f"Required package missing: {e.name}. Run in OSGeo4W Shell: "
            f"python -m pip install pypdf python-docx openpyxl pandas"
        )
    except Exception as e:
        return None, f"Error: {e}"


class ChatBridge(QObject):
    receiveMessage = pyqtSignal(str)
    setStatus = pyqtSignal(str)
    fileAttached = pyqtSignal(str)
    apiKeySaved = pyqtSignal()

    def __init__(self, agent_provider, parent=None):
        super().__init__(parent)
        self._agent_provider = agent_provider

    @pyqtSlot(str)
    def sendMessage(self, text):
        thread = threading.Thread(target=self._run, args=(text,), daemon=True)
        thread.start()

    def _run(self, text):
        try:
            agent = self._agent_provider()
            if agent is None:
                self.receiveMessage.emit(
                    "**No API key configured.** Click ⚙ in the header to set your OpenRouter API key."
                )
                return

            client = getattr(agent, "client", None)
            if client is not None and hasattr(client, "set_status_callback"):
                client.set_status_callback(lambda msg: self.setStatus.emit(msg))

            response = agent.run(text)

            if client is not None and hasattr(client, "set_status_callback"):
                client.set_status_callback(None)

            self.receiveMessage.emit(response if response else "_(empty response)_")
        except Exception as e:
            traceback.print_exc()
            self.receiveMessage.emit(f"**Error:** {e}")

    @pyqtSlot()
    def openSettings(self):
        settings = QgsSettings()
        current = settings.value(SETTINGS_KEY, "") or ""
        echo_mode = getattr(QLineEdit, "EchoMode", QLineEdit).Password
        key, ok = QInputDialog.getText(
            None,
            "QGIS AI Agent — Settings",
            "OpenRouter API key:",
            echo_mode,
            current,
        )
        if ok:
            settings.setValue(SETTINGS_KEY, key)
            self.setStatus.emit("API key saved" if key else "API key cleared")
            if key:
                self.apiKeySaved.emit()

    @pyqtSlot(result=str)
    def getInitialState(self):
        settings = QgsSettings()
        key = settings.value(SETTINGS_KEY, "") or ""
        done = settings.value(ONBOARDING_KEY, False)
        if isinstance(done, str):
            done = done.strip().lower() in ("true", "1", "yes")
        if not key and not done:
            return "onboarding"
        return "welcome"

    @pyqtSlot(str)
    def openUrl(self, url):
        from qgis.PyQt.QtCore import QUrl as _QUrl
        from qgis.PyQt.QtGui import QDesktopServices
        QDesktopServices.openUrl(_QUrl(url))

    @pyqtSlot()
    def markOnboardingComplete(self):
        QgsSettings().setValue(ONBOARDING_KEY, True)

    @pyqtSlot()
    def attachFile(self):
        import os
        filters = (
            "Supported files (*.pdf *.docx *.png *.jpg *.jpeg *.csv *.xlsx);;"
            "PDF (*.pdf);;Word (*.docx);;Images (*.png *.jpg *.jpeg);;"
            "CSV (*.csv);;Excel (*.xlsx);;All files (*)"
        )
        path, _ = QFileDialog.getOpenFileName(None, "Attach file", "", filters)
        if not path:
            return

        name = os.path.basename(path)
        data, err = _read_attached_file(path)
        if err is not None:
            self.fileAttached.emit(f"{name} — error: {err}")
            return

        self.fileAttached.emit(name)
        self.setStatus.emit("Analyzing...")
        thread = threading.Thread(
            target=self._analyze_file, args=(name, data), daemon=True
        )
        thread.start()

    def _analyze_file(self, name, data):
        try:
            agent = self._agent_provider()
            if agent is None:
                self.receiveMessage.emit(
                    "**No API key configured.** Click ⚙ in the header to set your OpenRouter API key."
                )
                return

            client = getattr(agent, "client", None)
            if client is not None and hasattr(client, "set_status_callback"):
                client.set_status_callback(lambda msg: self.setStatus.emit(msg))

            if data.get("is_image"):
                response = self._analyze_image(agent, name, data)
            else:
                prompt = (
                    f"I have attached a file: {name}\n\n"
                    f"File content:\n{data.get('text') or ''}\n\n"
                    "Please analyze this file and suggest what can be done with it in QGIS."
                )
                response = agent.run(prompt)

            if client is not None and hasattr(client, "set_status_callback"):
                client.set_status_callback(None)

            self.receiveMessage.emit(response if response else "_(empty response)_")
        except Exception as e:
            traceback.print_exc()
            self.receiveMessage.emit(f"**Error analyzing file:** {e}")
        finally:
            self.setStatus.emit("")

    def _analyze_image(self, agent, name, data):
        b64 = data.get("b64", "")
        mime = data.get("mime", "png")
        data_url = f"data:image/{mime};base64,{b64}"

        messages = [
            {
                "role": "system",
                "content": (
                    "You are a QGIS spatial analysis assistant. The user has attached an image; "
                    "describe what it shows and suggest how it could be used inside QGIS."
                ),
            }
        ]
        messages.extend(getattr(agent, "conversation_history", []))
        messages.append(
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            f"I have attached an image: {name}. "
                            "Please analyze it and suggest what can be done with it in QGIS."
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        )

        client = getattr(agent, "client", None)
        if client is None:
            return "Agent has no API client configured."

        result = client.complete(messages)
        if not isinstance(result, dict):
            return "[API error] Unexpected response from model client."
        if "error" in result:
            return f"[API error] {result['error']}"
        message = result.get("message") or {}
        content = message.get("content")
        if not content:
            return (
                "The current model did not return any analysis. "
                "Vision support depends on the model; try a vision-capable model in Settings."
            )
        try:
            agent.conversation_history.append(
                {"role": "user", "content": f"[Attached image: {name}]"}
            )
            agent.conversation_history.append({"role": "assistant", "content": content})
            if hasattr(agent, "_trim_history"):
                agent._trim_history()
        except Exception:
            pass
        return content


class QgisAiAgentDockWidget(QDockWidget):
    def __init__(self, agent_provider=None, parent=None):
        print("[QgisAiAgentDock] __init__")
        try:
            super().__init__("QGIS AI Agent", parent)
            self.setObjectName("QgisAiAgentDockWidget")
            self.setAllowedAreas(_dock_areas_left_right())

            container = QWidget(self)
            layout = QVBoxLayout(container)
            layout.setContentsMargins(0, 0, 0, 0)
            layout.setSpacing(0)

            self.web_view = None
            self.channel = None
            self.bridge = None

            if WEBENGINE_AVAILABLE:
                try:
                    self.web_view = QWebEngineView(container)
                    self.channel = QWebChannel(self.web_view.page())
                    self.bridge = ChatBridge(
                        agent_provider or (lambda: None), parent=self
                    )
                    self.channel.registerObject("bridge", self.bridge)
                    self.web_view.page().setWebChannel(self.channel)
                    self.web_view.setHtml(CHAT_HTML, QUrl("qrc:/qgis_ai_agent/"))
                    layout.addWidget(self.web_view)
                    print("[QgisAiAgentDock] QWebEngineView + QWebChannel ready")
                except Exception as e:
                    print(f"[QgisAiAgentDock] WebEngine init failed: {e}")
                    traceback.print_exc()
                    self.web_view = None

            if self.web_view is None:
                msg = "QtWebEngine is not available in this QGIS build."
                if _WEBENGINE_IMPORT_ERROR is not None:
                    msg += f"\n\nImport error: {_WEBENGINE_IMPORT_ERROR}"
                label = QLabel(msg, container)
                label.setWordWrap(True)
                label.setAlignment(_align_center())
                layout.addWidget(label)

            self.setWidget(container)
            print("[QgisAiAgentDock] init complete")
        except Exception as e:
            print(f"[QgisAiAgentDock] init crashed: {e}")
            traceback.print_exc()
            raise
